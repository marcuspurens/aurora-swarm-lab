"""Document extraction for PDF/DOCX/TXT."""

from __future__ import annotations

import os
import shutil
from pathlib import Path
from typing import List

from app.core.textnorm import normalize_whitespace


def _extract_txt(path: Path) -> str:
    return path.read_text(encoding="utf-8", errors="ignore")


def _extract_docx(path: Path) -> str:
    try:
        import docx  # type: ignore
    except Exception as exc:
        raise RuntimeError("python-docx not installed") from exc
    document = docx.Document(str(path))
    parts: List[str] = []
    for para in document.paragraphs:
        if para.text:
            parts.append(para.text)
    return "\n".join(parts)


def _extract_pdf_text_layer(path: Path) -> str:
    try:
        import pypdfium2 as pdfium  # type: ignore
    except Exception as exc:
        raise RuntimeError("pypdfium2 not installed") from exc

    doc = pdfium.PdfDocument(str(path))
    parts: List[str] = []
    try:
        for page_index in range(len(doc)):
            page = doc[page_index]
            textpage = page.get_textpage()
            try:
                parts.append(textpage.get_text_range())
            finally:
                textpage.close()
                page.close()
    finally:
        doc.close()
    return "\n".join(parts)


def _getenv_bool(name: str, default: bool) -> bool:
    raw = os.getenv(name)
    if raw is None:
        return default
    return str(raw).strip().lower() in {"1", "true", "yes", "on"}


def _getenv_int(name: str, default: int) -> int:
    raw = os.getenv(name)
    if raw is None:
        return default
    try:
        return int(str(raw).strip())
    except Exception:
        return default


def _getenv_float(name: str, default: float) -> float:
    raw = os.getenv(name)
    if raw is None:
        return default
    try:
        return float(str(raw).strip())
    except Exception:
        return default


def _pdf_ocr_enabled() -> bool:
    return _getenv_bool("AURORA_DOC_OCR_ENABLED", True)


def _pdf_ocr_min_text_chars() -> int:
    return max(0, _getenv_int("AURORA_DOC_OCR_MIN_TEXT_CHARS", 500))


def _pdf_ocr_max_pages() -> int:
    return max(1, _getenv_int("AURORA_DOC_OCR_MAX_PAGES", 120))


def _pdf_ocr_render_scale() -> float:
    scale = _getenv_float("AURORA_DOC_OCR_RENDER_SCALE", 2.0)
    return max(1.0, min(4.0, scale))


def _pdf_ocr_lang() -> str:
    value = str(os.getenv("AURORA_DOC_OCR_LANG", "eng") or "eng").strip()
    return value or "eng"


def _pdf_ocr_paddle_lang() -> str:
    configured = str(os.getenv("AURORA_DOC_OCR_PADDLE_LANG", "") or "").strip().lower()
    if configured:
        return configured
    tesseract_lang = _pdf_ocr_lang().strip().lower()
    if tesseract_lang in {"eng", "en"}:
        return "en"
    if tesseract_lang in {"swe", "sv"}:
        return "sv"
    return "en"


def _pdf_ocr_backend_order() -> List[str]:
    raw = str(os.getenv("AURORA_DOC_OCR_BACKEND", "auto") or "auto").strip().lower()
    if raw == "paddleocr":
        return ["paddleocr", "tesseract"]
    if raw == "tesseract":
        return ["tesseract", "paddleocr"]
    return ["paddleocr", "tesseract"]


def _should_attempt_pdf_ocr(text: str) -> bool:
    if not _pdf_ocr_enabled():
        return False
    text_chars = len(normalize_whitespace(text or ""))
    return text_chars < _pdf_ocr_min_text_chars()


def _build_paddle_ocr_reader():
    try:
        from paddleocr import PaddleOCR  # type: ignore
    except Exception:
        return None
    try:
        return PaddleOCR(use_angle_cls=True, lang=_pdf_ocr_paddle_lang())
    except Exception:
        return None


def _extract_text_from_paddle_result(result: object) -> str:
    out: List[str] = []
    stack: List[object] = [result]
    while stack:
        current = stack.pop()
        if not isinstance(current, (list, tuple)):
            continue
        if len(current) >= 2 and isinstance(current[1], (list, tuple)) and current[1]:
            candidate = current[1][0]
            if isinstance(candidate, str):
                text = candidate.strip()
                if text:
                    out.append(text)
                continue
        for item in reversed(current):
            stack.append(item)
    return "\n".join(out)


def _ocr_image_with_paddle(image: object, reader: object) -> str:
    try:
        import numpy as np  # type: ignore
    except Exception:
        return ""
    try:
        result = reader.ocr(np.array(image), cls=True)
    except Exception:
        return ""
    return _extract_text_from_paddle_result(result)


def _ocr_image_with_tesseract(image: object) -> str:
    if shutil.which("tesseract") is None:
        return ""
    try:
        import pytesseract  # type: ignore
    except Exception:
        return ""
    try:
        text = str(pytesseract.image_to_string(image, lang=_pdf_ocr_lang()) or "").strip()
    except Exception:
        return ""
    return text


def _extract_pdf_ocr(path: Path) -> str:
    try:
        import pypdfium2 as pdfium  # type: ignore
    except Exception:
        return ""

    parts: List[str] = []
    doc = pdfium.PdfDocument(str(path))
    max_pages = _pdf_ocr_max_pages()
    scale = _pdf_ocr_render_scale()
    backend_order = _pdf_ocr_backend_order()
    paddle_reader = _build_paddle_ocr_reader() if "paddleocr" in backend_order else None
    try:
        page_count = min(len(doc), max_pages)
        for page_index in range(page_count):
            page = doc[page_index]
            bitmap = None
            image = None
            try:
                bitmap = page.render(scale=scale)
                image = bitmap.to_pil()
                text = ""
                for backend in backend_order:
                    if backend == "paddleocr":
                        if paddle_reader is None:
                            continue
                        text = _ocr_image_with_paddle(image, paddle_reader)
                    elif backend == "tesseract":
                        text = _ocr_image_with_tesseract(image)
                    if text:
                        break
                if text:
                    parts.append(text)
            except Exception:
                # Keep extraction resilient; OCR can fail per page.
                continue
            finally:
                try:
                    if image is not None and hasattr(image, "close"):
                        image.close()
                except Exception:
                    pass
                try:
                    if bitmap is not None and hasattr(bitmap, "close"):
                        bitmap.close()
                except Exception:
                    pass
                page.close()
    finally:
        doc.close()
    return "\n".join(parts)


def _extract_pdf(path: Path) -> str:
    text_layer = _extract_pdf_text_layer(path)
    if not _should_attempt_pdf_ocr(text_layer):
        return text_layer
    ocr_text = _extract_pdf_ocr(path)
    if not ocr_text.strip():
        return text_layer
    if len(normalize_whitespace(ocr_text)) > len(normalize_whitespace(text_layer)):
        return ocr_text
    return text_layer


def extract(path: str) -> str:
    file_path = Path(path)
    suffix = file_path.suffix.lower()
    if suffix in {".txt", ".md"}:
        text = _extract_txt(file_path)
    elif suffix == ".docx":
        text = _extract_docx(file_path)
    elif suffix == ".pdf":
        text = _extract_pdf(file_path)
    else:
        raise ValueError(f"Unsupported document type: {suffix}")
    return normalize_whitespace(text)
