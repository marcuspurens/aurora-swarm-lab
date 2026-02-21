from pathlib import Path

from app.modules.doc_extract import extract_doc
from app.modules.doc_extract.extract_doc import extract


def test_extract_txt_normalizes(tmp_path):
    path = tmp_path / "sample.txt"
    path.write_text("Hello\n\nworld", encoding="utf-8")
    text = extract(str(path))
    assert text == "Hello world"


def test_extract_docx(tmp_path):
    try:
        import docx  # type: ignore
    except Exception:  # pragma: no cover - environment issue
        return

    path = tmp_path / "sample.docx"
    doc = docx.Document()
    doc.add_paragraph("First")
    doc.add_paragraph("Second")
    doc.save(str(path))

    text = extract(str(path))
    assert "First" in text and "Second" in text


def test_extract_pdf_uses_ocr_when_text_layer_is_thin(tmp_path, monkeypatch):
    path = tmp_path / "scan.pdf"
    path.write_bytes(b"%PDF-1.4\n")

    monkeypatch.setenv("AURORA_DOC_OCR_ENABLED", "1")
    monkeypatch.setenv("AURORA_DOC_OCR_MIN_TEXT_CHARS", "50")
    monkeypatch.setattr(extract_doc, "_extract_pdf_text_layer", lambda _path: "x")
    monkeypatch.setattr(extract_doc, "_extract_pdf_ocr", lambda _path: "Scanned OCR text")

    text = extract(str(path))
    assert text == "Scanned OCR text"


def test_extract_pdf_skips_ocr_when_disabled(tmp_path, monkeypatch):
    path = tmp_path / "scan.pdf"
    path.write_bytes(b"%PDF-1.4\n")

    monkeypatch.setenv("AURORA_DOC_OCR_ENABLED", "0")
    monkeypatch.setattr(extract_doc, "_extract_pdf_text_layer", lambda _path: "Text layer value")
    monkeypatch.setattr(extract_doc, "_extract_pdf_ocr", lambda _path: "Scanned OCR text")

    text = extract(str(path))
    assert text == "Text layer value"


def test_pdf_ocr_backend_order_auto_prefers_paddle(monkeypatch):
    monkeypatch.delenv("AURORA_DOC_OCR_BACKEND", raising=False)
    assert extract_doc._pdf_ocr_backend_order() == ["paddleocr", "tesseract"]


def test_pdf_ocr_backend_order_tesseract_first(monkeypatch):
    monkeypatch.setenv("AURORA_DOC_OCR_BACKEND", "tesseract")
    assert extract_doc._pdf_ocr_backend_order() == ["tesseract", "paddleocr"]
